"""
Document loader and chunker utility for RAG system.
Supports PDF, DOCX, and TXT files.
"""

import os
from typing import List, Dict, Any
from pathlib import Path
import PyPDF2
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document


class DocumentLoader:
    """Load and chunk documents for RAG"""

    def __init__(self, documents_path: str = "../Documents", chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize document loader

        Args:
            documents_path: Path to documents directory
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        self.documents_path = Path(documents_path)
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Text splitter for chunking
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", " ", ""],
        )

    def load_pdf(self, file_path: str) -> str:
        """
        Load text from PDF file

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text from PDF
        """
        text = ""
        try:
            with open(file_path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += f"\n--- Page {page_num + 1} ---\n"
                    text += page.extract_text()
            return text
        except Exception as e:
            raise Exception(f"Error loading PDF {file_path}: {str(e)}")

    def load_txt(self, file_path: str) -> str:
        """Load text from TXT file"""
        try:
            with open(file_path, "r", encoding="utf-8") as txt_file:
                return txt_file.read()
        except Exception as e:
            raise Exception(f"Error loading TXT {file_path}: {str(e)}")

    def load_docx(self, file_path: str) -> str:
        """Load text from DOCX file"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error loading DOCX {file_path}: {str(e)}")

    def load_document_by_name(self, filename: str) -> str:
        """
        Load document by filename

        Args:
            filename: Name of the file (e.g., "Investment_Insights.pdf")

        Returns:
            Extracted text from document
        """
        file_path = self.documents_path / filename

        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {filename}")

        if filename.endswith(".pdf"):
            return self.load_pdf(str(file_path))
        elif filename.endswith(".txt"):
            return self.load_txt(str(file_path))
        elif filename.endswith(".docx"):
            return self.load_docx(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {filename}")

    def chunk_document(self, text: str, metadata: Dict[str, Any] = None) -> List[Document]:
        """
        Chunk text into smaller pieces

        Args:
            text: Text to chunk
            metadata: Metadata to attach to chunks

        Returns:
            List of LangChain Document objects
        """
        chunks = self.text_splitter.split_text(text)
        documents = []
        
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    **(metadata or {}),
                    "chunk_id": i,
                    "chunk_size": len(chunk),
                }
            )
            documents.append(doc)
        
        return documents

    def load_and_chunk_document(
        self,
        filename: str,
        metadata: Dict[str, Any] = None,
    ) -> List[Document]:
        """
        Load and chunk a document in one step

        Args:
            filename: Name of the file
            metadata: Additional metadata

        Returns:
            List of chunked documents
        """
        text = self.load_document_by_name(filename)
        metadata = metadata or {}
        metadata["source"] = filename
        return self.chunk_document(text, metadata)

    def load_all_documents_in_directory(self) -> Dict[str, List[Document]]:
        """
        Load and chunk all documents in the directory

        Returns:
            Dictionary with filename as key and list of chunks as value
        """
        results = {}
        
        if not self.documents_path.exists():
            raise FileNotFoundError(f"Documents directory not found: {self.documents_path}")
        
        for file_path in self.documents_path.iterdir():
            if file_path.is_file() and file_path.suffix in [".pdf", ".txt", ".docx"]:
                try:
                    chunks = self.load_and_chunk_document(file_path.name)
                    results[file_path.name] = chunks
                    print(f"✓ Loaded {file_path.name}: {len(chunks)} chunks")
                except Exception as e:
                    print(f"✗ Error loading {file_path.name}: {str(e)}")
        
        return results
